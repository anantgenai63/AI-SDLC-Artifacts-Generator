# =========================================================
# ENTERPRISE AI SDLC GENERATOR
# Author: Anantkreshna Vedanarayanan
# Chennai, India
# =========================================================

import streamlit as st
from openai import OpenAI
from docx import Document
import io
import zipfile
from datetime import datetime

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="Enterprise AI SDLC Generator",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================================================
# LLM CONFIG
# =========================================================

GROQ_API_KEY = "API-KEY-HERE" # Replace with your actual API key

client = OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1"
)

MODEL_NAME = "llama-3.1-8b-instant"
#MODEL_NAME = "llama-3.3-70b-versatile"

# =========================================================
# SESSION STATE
# =========================================================

if "results" not in st.session_state:
    st.session_state.results = {}

# =========================================================
# CUSTOM CSS
# =========================================================

def inject_css():

    st.markdown("""
    <style>

    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }

    .stApp {
        background: linear-gradient(180deg,#f8fafc 0%,#eef2ff 100%);
    }

    .topbar {
        background: linear-gradient(90deg,#4338ca,#7c3aed);
        padding: 1.5rem 2rem;
        border-radius: 0 0 18px 18px;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    .topbar h1 {
        margin: 0;
        font-size: 1.8rem;
        font-weight: 700;
    }

    .topbar p {
        margin-top: 0.3rem;
        opacity: 0.9;
    }

    .artifact-card {
        background: white;
        border-radius: 16px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        border: 1px solid #e5e7eb;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }

    .metric-box {
        background: white;
        padding: 1rem;
        border-radius: 14px;
        text-align: center;
        border: 1px solid #e5e7eb;
    }

    .stButton > button {
        width: 100%;
        background: linear-gradient(90deg,#4f46e5,#7c3aed);
        color: white;
        border-radius: 10px;
        border: none;
        font-weight: 600;
        height: 3rem;
    }

    .stButton > button:hover {
        transform: translateY(-1px);
    }

    section[data-testid="stSidebar"] {
        background: #ffffff;
        border-right: 1px solid #e5e7eb;
    }

    </style>
    """, unsafe_allow_html=True)

inject_css()

# =========================================================
# HEADER
# =========================================================

st.markdown("""
<div class="topbar">
    <h1>🚀 Enterprise AI SDLC Generator</h1>
    <p>AI-powered multi-agent SDLC document generation platform</p>
</div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:

    st.markdown("## ⚙️ Configuration")

    project_name = st.text_input(
        "Project Name",
        placeholder="AI Approval Workflow"
    )

    industry = st.selectbox(
        "Industry",
        [
            "Banking",
            "Insurance",
            "Healthcare",
            "Retail",
            "Telecom",
            "Manufacturing"
        ]
    )

    project_type = st.selectbox(
        "Project Type",
        [
            "Workflow Automation",
            "AI Platform",
            "Enterprise Application",
            "Data Platform",
            "Web Application"
        ]
    )

    compliance = st.multiselect(
        "Compliance Requirements",
        [
            "GDPR",
            "SOC2",
            "ISO27001",
            "HIPAA",
            "PCI-DSS",
            "SOX"
        ]
    )

    st.markdown("## 📦 Deliverables")

    generate_frd = st.checkbox("FRD", value=True)
    generate_user_stories = st.checkbox("User Stories", value=True)
    generate_architecture = st.checkbox("Architecture", value=True)
    generate_database = st.checkbox("Database Design", value=True)
    generate_test_cases = st.checkbox("Test Cases", value=True)
    generate_traceability = st.checkbox("Traceability Matrix", value=True)
    generate_manual = st.checkbox("User Manual", value=True)
    generate_yaml = st.checkbox("YAML / DSL", value=True)
    generate_code = st.checkbox("Starter Code", value=True)
    generate_ui_wireframe = st.checkbox("UI Wireframes", value=True)

# =========================================================
# MAIN INPUT
# =========================================================

st.markdown("## 🧠 Business Use Case")

use_case = st.text_area(
    "Describe your business problem",
    height=280,
    placeholder="""
Example:

Build an AI-assisted service deviation and approval workflow platform.

Features:
- Approval workflow
- AI risk scoring
- SLA management
- Notifications
- Compliance validation
- Audit trail
- Reporting dashboard
- Role-based access control
"""
)

# =========================================================
# AI GENERATOR
# =========================================================

def generate_artifact(prompt):

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {
                "role": "system",
                "content": """
You are a world-class Enterprise Architect,
Business Analyst,
Solution Designer,
Security Architect,
UI/UX Designer,
QA Architect,
and AI Transformation Expert.

Generate enterprise-grade SDLC deliverables.

Every output must:
- Be professionally structured
- Be enterprise-ready
- Include governance
- Include security
- Include compliance
- Include scalability
- Include auditability
- Include production-level detail

Where applicable:
- Generate Mermaid diagrams
- Generate UI wireframes
- Generate workflow diagrams
- Generate architecture layers
- Generate integration diagrams
"""
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.2,
        max_tokens=4000
    )

    return response.choices[0].message.content

# =========================================================
# PROMPTS
# =========================================================

PROMPTS = {

    "FRD": """
Generate enterprise Functional Requirements Document.

Include:
- Executive Summary
- Objectives
- Scope
- Actors and Roles
- Workflow States
- Workflow Transitions
- Functional Requirements
- Non Functional Requirements
- SLA
- Security
- Compliance
- Audit
- Escalation
- Notifications
- Exception Handling
- Validation Rules
- Acceptance Criteria
""",

    "USER_STORIES": """
Generate Agile User Stories.

Include:
- Epic
- User Story
- Acceptance Criteria
- Priority
- Dependencies
- Definition of Done
""",

    "ARCHITECTURE": """
Generate Enterprise Architecture Document.

Include:
- Application Architecture
- Technical Architecture
- Integration Architecture
- Security Architecture
- Deployment Architecture

Generate:
- Mermaid diagrams
- Component diagrams
- Sequence diagrams
- Integration flow diagrams
""",

    "DATABASE": """
Generate Enterprise Database Design.

Include:
- Logical Data Model
- Physical Data Model
- ER Diagram
- PostgreSQL schema
- Relationships
- Constraints
- Index strategy
""",

    "TEST_CASES": """
Generate Enterprise Functional Test Cases.

Include:
- Positive testing
- Negative testing
- Integration testing
- Security testing
- UAT scenarios

Output in tabular format.
""",

    "TRACEABILITY": """
Generate Requirement Traceability Matrix.

Include:
- Requirement ID
- Requirement
- User Story
- Test Case Mapping
- Status
""",

    "MANUAL": """
Generate End User Manual.

Include:
- Introduction
- User Navigation
- Screen Descriptions
- Workflow Guide
- FAQs
- Troubleshooting
""",

    "YAML": """
Generate enterprise YAML / DSL specification.

Include:
- Workflow states
- Transitions
- Roles
- Rules
- APIs
- Events
""",

    "CODE": """
Generate enterprise starter code architecture.

Include:
- Streamlit frontend
- FastAPI backend
- PostgreSQL integration
- Authentication
- AI orchestration
- Folder structure
- Deployment structure
""",

    "UI_WIREFRAME": """
Generate enterprise UI wireframes.

Include:
- Login screen
- Dashboard
- Approval workflow screen
- Reports dashboard
- Notifications panel
- Admin screen

Generate:
- ASCII wireframes
- Mermaid UI diagrams
"""
}

# =========================================================
# DOCX CREATOR
# =========================================================

def create_docx(title, content):

    doc = Document()

    doc.add_heading(title, level=1)

    doc.add_paragraph(content)

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()

# =========================================================
# ZIP CREATOR
# =========================================================

def create_zip(files):

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "a") as zip_file:

        for filename, data in files.items():

            zip_file.writestr(filename, data)

    return zip_buffer.getvalue()

# =========================================================
# MULTI AGENT EXECUTION
# =========================================================

AGENTS = {}

if generate_frd:
    AGENTS["FRD"] = "FRD.docx"

if generate_user_stories:
    AGENTS["USER_STORIES"] = "UserStories.docx"

if generate_architecture:
    AGENTS["ARCHITECTURE"] = "Architecture.docx"

if generate_database:
    AGENTS["DATABASE"] = "DatabaseDesign.docx"

if generate_test_cases:
    AGENTS["TEST_CASES"] = "TestCases.docx"

if generate_traceability:
    AGENTS["TRACEABILITY"] = "TraceabilityMatrix.docx"

if generate_manual:
    AGENTS["MANUAL"] = "UserManual.docx"

if generate_yaml:
    AGENTS["YAML"] = "workflow.yaml"

if generate_code:
    AGENTS["CODE"] = "StarterCode.docx"

if generate_ui_wireframe:
    AGENTS["UI_WIREFRAME"] = "UIWireframes.docx"

# =========================================================
# GENERATE BUTTON
# =========================================================

if st.button("🚀 Generate Enterprise Deliverables"):

    if not use_case.strip():

        st.warning("Please enter business use case")

        st.stop()

    progress = st.progress(0)

    results = {}

    export_files = {}

    total = len(AGENTS)

    for idx, (artifact, filename) in enumerate(AGENTS.items()):

        with st.spinner(f"Generating {artifact}..."):

            prompt = f"""
Project Name:
{project_name}

Industry:
{industry}

Project Type:
{project_type}

Compliance:
{', '.join(compliance)}

Use Case:
{use_case}

Task:
{PROMPTS[artifact]}
"""

            output = generate_artifact(prompt)

            results[artifact] = output

            if filename.endswith(".yaml"):

                export_files[filename] = output

            else:

                export_files[filename] = create_docx(
                    artifact,
                    output
                )

        progress.progress((idx + 1) / total)

    st.session_state.results = results

    st.success("✅ Enterprise Deliverables Generated Successfully")

# =========================================================
# RESULTS
# =========================================================

if st.session_state.results:

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(f"""
        <div class="metric-box">
            <h2>{len(st.session_state.results)}</h2>
            <p>Artifacts Generated</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="metric-box">
            <h2>{industry}</h2>
            <p>Industry</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown(f"""
        <div class="metric-box">
            <h2>{project_type}</h2>
            <p>Project Type</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    for key, value in st.session_state.results.items():

        with st.expander(f"📄 {key}", expanded=False):

            st.markdown(value)

    # =====================================================
    # ZIP DOWNLOAD
    # =====================================================

    files = {}

    for key, value in st.session_state.results.items():

        if key == "YAML":

            files["workflow.yaml"] = value

        else:

            files[f"{key}.docx"] = create_docx(
                key,
                value
            )

    zip_data = create_zip(files)

    st.download_button(
        label="📥 Download All Deliverables",
        data=zip_data,
        file_name=f"{project_name}_deliverables.zip",
        mime="application/zip"
    )

# =========================================================
# EMPTY STATE
# =========================================================

else:

    st.markdown("""
    <div class="artifact-card">

    <h3>✨ Enterprise AI SDLC Platform</h3>

    This platform generates:

    - Functional Requirements Documents
    - User Stories
    - Technical Architecture
    - Database Design
    - UI Wireframes
    - Test Cases
    - Traceability Matrix
    - YAML / DSL Specs
    - User Manuals
    - Starter Code


    </div>
    """, unsafe_allow_html=True)
